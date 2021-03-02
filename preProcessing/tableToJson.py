import docx
import json

# constant
JSON_FILE = "./participants.json"
DOC_NAME = "./participants.docx"

# contains all the data to be formatted
DATA = {}
# to check used groups, or check repetition of group names
GROUPS = {}

def main ():
	iteration = 0

	print("stat")

	document = docx.Document(DOC_NAME)
	iteration = initialise_data_for_grade(document,"XII", iteration)
	iteration = initialise_data_for_grade(document,"XI", iteration)
	iteration = initialise_data_for_grade(document,"X", iteration)
	iteration = initialise_data_for_grade(document,"IX", iteration)

	jsonified = json.dumps(DATA)

	file = open(JSON_FILE, "w")
	file.write(jsonified)
	file.flush()
	file.close()


def initialise_data_for_grade(document,grade,iteration):
	table = document.tables[iteration]
	sub_map = {}

	for  i in range(0,4):
		cells = table.column_cells(i)
		group_name = cells[0].paragraphs[0].text
		participants = []

		if GROUPS.get(group_name) != None:
			participants = sub_map[group_name] 
		elif group_name != "":
			GROUPS[group_name] = "helo"

		for j in range(1,len(cells)):
			name = cells[j].paragraphs[0].text
			if name != "":
				# print(name)
				participants.append(name)

		sub_map[group_name] = participants

	DATA[grade] = sub_map

	return iteration + 1

main()